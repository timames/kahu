"""Generate Kahu pricing one-sheet as .docx"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# -- Page setup: letter, narrow margins --
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)

BRAND_DARK = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_ACCENT = RGBColor(0x00, 0x7A, 0x5E)
BRAND_GOLD = RGBColor(0xC8, 0x96, 0x2E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x66, 0x66, 0x66)


def add_colored_heading(text, color, size=20, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.name = "Calibri"
    return p


def add_section_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_ACCENT
    run.font.bold = True
    run.font.name = "Calibri"
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="007A5E"/>'
        f'</w:pBdr>'
    )
    p._p.get_or_add_pPr().append(pBdr)
    return p


def add_body(text, bold=False, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def shade_cell(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="DDDDDD"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="2" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="2" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="2" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="2" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def format_cell(cell, text, size=9, bold=False, color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    run.font.color.rgb = color or RGBColor(0x33, 0x33, 0x33)


# ============================================================
# HEADER
# ============================================================
add_colored_heading("KAHU", BRAND_DARK, size=28)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(2)
run = p.add_run("Pricing & Plans")
run.font.size = Pt(14)
run.font.color.rgb = BRAND_ACCENT
run.font.bold = True
run.font.name = "Calibri"

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run("by ComplyHI  |  complyhi.com  |  Honolulu, HI")
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.font.italic = True
run.font.name = "Calibri"

# Separator
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(4)
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="1A1A2E"/>'
    f'</w:pBdr>'
)
p._p.get_or_add_pPr().append(pBdr)

# ============================================================
# TAGLINE
# ============================================================
add_body(
    "Enterprise-grade AI security operations for the price of a software subscription. "
    "Every plan includes the full triage pipeline, mobile PWA, gamified analyst experience, "
    "and on-premises AI \u2014 because your security shouldn't depend on your budget.",
    size=9.5,
)

# ============================================================
# PRICING TABLE
# ============================================================
add_section_header("Plans")

# 4 columns: Feature | Solo | Squad | Command
table = doc.add_table(rows=16, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

# Column headers
headers = ["", "Solo", "Squad", "Command"]
subtitles = ["", "Small office", "Mid-size business", "Enterprise / MSP"]
prices_hw = ["", "$499 one-time", "$2,499 one-time", "Virtual appliance"]
prices_mo = ["", "$49/mo", "$149/mo", "$299/mo per 150 endpoints"]

for col_idx, header in enumerate(headers):
    cell = table.cell(0, col_idx)
    if col_idx > 0:
        shade_cell(cell, "1A1A2E")
        format_cell(cell, header, size=12, bold=True, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_borders(cell, "1A1A2E")

for col_idx, sub in enumerate(subtitles):
    cell = table.cell(1, col_idx)
    if col_idx > 0:
        shade_cell(cell, "1A1A2E")
        format_cell(cell, sub, size=8, color=RGBColor(0xAA, 0xAA, 0xAA), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_borders(cell, "1A1A2E")

for col_idx, price in enumerate(prices_hw):
    cell = table.cell(2, col_idx)
    if col_idx == 0:
        format_cell(cell, "Hardware", size=9, bold=True, color=BRAND_DARK)
    else:
        shade_cell(cell, "F0F4F3")
        format_cell(cell, price, size=9, bold=True, color=BRAND_DARK, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_borders(cell)

for col_idx, price in enumerate(prices_mo):
    cell = table.cell(3, col_idx)
    if col_idx == 0:
        format_cell(cell, "Subscription", size=9, bold=True, color=BRAND_DARK)
    else:
        shade_cell(cell, "F0F4F3")
        format_cell(cell, price, size=9, bold=True, color=BRAND_ACCENT, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_borders(cell)

# Feature rows
features = [
    ("Endpoints", "\u226425", "25\u2013150", "Unlimited (per block)"),
    ("Analyst seats", "1", "5", "Unlimited"),
    ("AI Triage Pipeline", "\u2713", "\u2713", "\u2713"),
    ("Mobile PWA + Swipe UX", "\u2713", "\u2713", "\u2713"),
    ("Gamification & Coach", "\u2713", "\u2713", "\u2713"),
    ("NL Investigation (Ask)", "\u2713", "\u2713", "\u2713"),
    ("On-premises LLM", "CPU", "GPU-accelerated", "GPU-accelerated"),
    ("Compliance frameworks", "1", "3", "All"),
    ("Threat intel updates", "Community", "Managed", "Managed + custom rules"),
    ("Multi-site dashboard", "\u2014", "Up to 5 sites", "Unlimited"),
    ("API / SOAR integration", "\u2014", "\u2014", "\u2713"),
    ("Hardware", "NAS chassis", "1U server + GPU", "Your infrastructure"),
]

for row_idx, (feature, solo, squad, command) in enumerate(features):
    actual_row = row_idx + 4
    cell0 = table.cell(actual_row, 0)
    cell1 = table.cell(actual_row, 1)
    cell2 = table.cell(actual_row, 2)
    cell3 = table.cell(actual_row, 3)

    format_cell(cell0, feature, size=9, bold=True, color=BRAND_DARK)
    format_cell(cell1, solo, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    format_cell(cell2, squad, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    format_cell(cell3, command, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Alternate row shading
    if row_idx % 2 == 0:
        for c in [cell0, cell1, cell2, cell3]:
            shade_cell(c, "FAFAFA")

    for c in [cell0, cell1, cell2, cell3]:
        set_cell_borders(c)

# ============================================================
# THE COMPARISON
# ============================================================
add_section_header("How Kahu Compares")

comp_table = doc.add_table(rows=5, cols=3)
comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER

comp_headers = ["Alternative", "Typical monthly cost", "What you get"]
for col_idx, h in enumerate(comp_headers):
    cell = comp_table.cell(0, col_idx)
    shade_cell(cell, "1A1A2E")
    format_cell(cell, h, size=9, bold=True, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_borders(cell, "1A1A2E")

comparisons = [
    ("MSSP contract", "$2,000\u2013$10,000/mo", "Outsourced monitoring, shared analysts, data leaves premises"),
    ("Cloud SIEM + analyst", "$5,000\u2013$15,000/mo", "Splunk/Sentinel license + FTE salary, cloud dependency"),
    ("Kahu Solo", "$49/mo + $499 box", "Full AI SOC, on-prem, mobile-first, compliance evidence"),
    ("Kahu Squad", "$149/mo + $2,499 box", "Everything above + GPU inference, 5 seats, multi-site"),
]

for row_idx, (alt, cost, desc) in enumerate(comparisons):
    actual_row = row_idx + 1
    format_cell(comp_table.cell(actual_row, 0), alt, size=9, bold=True, color=BRAND_DARK)
    format_cell(comp_table.cell(actual_row, 1), cost, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    format_cell(comp_table.cell(actual_row, 2), desc, size=8.5)

    # Highlight Kahu rows
    if "Kahu" in alt:
        for c_idx in range(3):
            shade_cell(comp_table.cell(actual_row, c_idx), "E8F5E9")

    for c_idx in range(3):
        set_cell_borders(comp_table.cell(actual_row, c_idx))

# ============================================================
# WHAT EVERY PLAN INCLUDES
# ============================================================
add_section_header("Every Plan Includes")

includes = [
    ("Full AI triage pipeline",
     "4-stage alert processing (filter, enrich, LLM triage, human disposition). "
     "AI advises, you decide. Works in degraded mode if AI is offline."),
    ("Mobile PWA with swipe-to-triage",
     "Manage security from your phone. Swipe alerts like cards \u2014 right to confirm, "
     "left to dismiss, up to escalate. Haptic feedback, offline queuing."),
    ("Gamification & coaching",
     "Score tracking, streaks, badges, rank progression (Recruit to Warden), "
     "and micro-lessons after every triage decision. Security becomes engaging."),
    ("100% on-premises AI",
     "No cloud. No telemetry. No data leaves your network. "
     "Meets ITAR, CUI, HIPAA, and data sovereignty requirements."),
    ("30-day free trial",
     "Try the virtual appliance free for 30 days. Full functionality, no credit card required."),
]

inc_table = doc.add_table(rows=len(includes), cols=2)
inc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (title, desc) in enumerate(includes):
    cell0 = inc_table.cell(i, 0)
    cell1 = inc_table.cell(i, 1)
    cell0.width = Inches(2.0)
    cell1.width = Inches(4.5)

    format_cell(cell0, title, size=9, bold=True, color=BRAND_DARK)
    shade_cell(cell0, "F0F4F3")
    format_cell(cell1, desc, size=9)

    for c in [cell0, cell1]:
        tc = c._tc
        tcPr = tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="nil"/>'
            f'  <w:left w:val="nil"/>'
            f'  <w:right w:val="nil"/>'
            f'  <w:bottom w:val="single" w:sz="2" w:space="0" w:color="DDDDDD"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)

# ============================================================
# CTA
# ============================================================
add_section_header("Get Started")
add_body(
    "Start your 30-day free trial today. Deploy the virtual appliance on your existing "
    "infrastructure and see AI-powered security operations in action \u2014 no hardware purchase required.",
    size=10,
    bold=True,
)

# ============================================================
# FOOTER
# ============================================================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="1A1A2E"/>'
    f'</w:pBdr>'
)
p._p.get_or_add_pPr().append(pBdr)
run = p.add_run("KAHU by ComplyHI")
run.font.size = Pt(9)
run.font.bold = True
run.font.color.rgb = BRAND_DARK
run.font.name = "Calibri"
run = p.add_run("  |  On-Premises AI Security Operations  |  complyhi.com  |  info@complyhi.com")
run.font.size = Pt(8)
run.font.color.rgb = GRAY
run.font.name = "Calibri"

# SAVE
out_path = os.path.join(os.path.dirname(__file__), "Kahu_Pricing_OneSheet.docx")
doc.save(out_path)
print(f"Saved to: {out_path}")
