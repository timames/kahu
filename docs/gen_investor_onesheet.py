"""Generate Kahu investor one-sheet as .docx"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# -- Page setup: letter, narrow margins --
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)

BRAND_DARK = RGBColor(0x1A, 0x1A, 0x2E)   # deep navy
BRAND_ACCENT = RGBColor(0x00, 0x7A, 0x5E)  # teal green
BRAND_GOLD = RGBColor(0xC8, 0x96, 0x2E)    # gold accent
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x66, 0x66, 0x66)

def add_colored_heading(text, color, size=20, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.name = "Calibri"
    return p

def add_section_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_ACCENT
    run.font.bold = True
    run.font.name = "Calibri"
    # Add a thin line
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="007A5E"/>'
        f'</w:pBdr>'
    )
    p._p.get_or_add_pPr().append(pBdr)
    return p

def add_body(text, bold=False, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.name = "Calibri"
        run.font.color.rgb = BRAND_DARK
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def shade_cell(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

# ============================================================
# HEADER
# ============================================================
add_colored_heading("KAHU", BRAND_DARK, size=28, alignment=WD_ALIGN_PARAGRAPH.LEFT)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(2)
run = p.add_run("On-Premises AI Security Operations Appliance")
run.font.size = Pt(13)
run.font.color.rgb = BRAND_ACCENT
run.font.bold = True
run.font.name = "Calibri"

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run("by ComplyHI  |  complyhi.com  |  Honolulu, HI")
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.font.italic = True
run.font.name = "Calibri"

# Thin separator
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(4)
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="1A1A2E"/>'
    f'</w:pBdr>'
)
p._p.get_or_add_pPr().append(pBdr)

# ============================================================
# THE PROBLEM
# ============================================================
add_section_header("The Problem")
add_body(
    "Defense contractors, healthcare providers, and critical infrastructure operators face "
    "a perfect storm: escalating cyber threats, strict compliance mandates (CMMC, NIST 800-171, HIPAA), "
    "and an acute shortage of trained security analysts. Current solutions force a painful tradeoff \u2014 "
    "send sensitive data to cloud-based AI platforms (violating data sovereignty requirements) or "
    "drown in manual alert triage and spreadsheet-driven compliance.",
    size=9.5
)

# ============================================================
# THE SOLUTION
# ============================================================
add_section_header("The Solution")
add_body(
    "Kahu is a turnkey security appliance that runs entirely on-premises. It combines a full "
    "SIEM/XDR platform (Wazuh) with a locally-hosted AI engine (Ollama) and proprietary orchestration "
    "that automates alert triage, investigation, and compliance evidence generation \u2014 with zero data "
    "leaving the customer's network.",
    size=9.5
)

# ============================================================
# KEY CAPABILITIES (2-column table)
# ============================================================
add_section_header("Key Capabilities")

capabilities = [
    ("AI-Powered Alert Triage",
     "4-stage pipeline (filter, enrich, LLM triage, human disposition) reduces analyst alert fatigue by 80%+. AI advises; humans decide."),
    ("Natural-Language Investigation",
     "Analysts ask questions in plain English \u2014 \"Show me failed logins from external IPs this week\" \u2014 and get instant, sourced answers."),
    ("Automated Compliance Evidence",
     "Every triage action auto-generates hash-chained, tamper-evident evidence mapped to NIST 800-171, CMMC L2, HIPAA, and CIS Controls."),
    ("100% On-Premises AI",
     "Local LLM inference via Ollama. No cloud fallback. No telemetry. Full data sovereignty from day one. Meets ITAR and CUI requirements."),
    ("Plug-and-Play Deployment",
     "Ships as a pre-configured appliance (hardware or virtual). Wizard-driven setup. SOC operational in under 4 hours, not weeks."),
    ("Vulnerability Scanning",
     "Built-in lightweight vulnerability detection integrated with the triage pipeline. No additional vendor licenses required."),
]

table = doc.add_table(rows=len(capabilities), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

for i, (title, desc) in enumerate(capabilities):
    cell0 = table.cell(i, 0)
    cell1 = table.cell(i, 1)
    cell0.width = Inches(1.8)
    cell1.width = Inches(4.7)

    # Title cell
    p = cell0.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = BRAND_DARK
    run.font.name = "Calibri"
    shade_cell(cell0, "F0F4F3")

    # Desc cell
    p = cell1.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(desc)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    run.font.name = "Calibri"

# Remove table borders, add subtle row lines
for row in table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="nil"/>'
            f'  <w:left w:val="nil"/>'
            f'  <w:right w:val="nil"/>'
            f'  <w:bottom w:val="single" w:sz="2" w:space="0" w:color="DDDDDD"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)

# ============================================================
# MARKET OPPORTUNITY
# ============================================================
add_section_header("Market Opportunity")
add_body(
    "The global SIEM market is $6.4B (2025) growing at 14% CAGR. The AI-in-cybersecurity market "
    "reaches $38B by 2028. CMMC 2.0 enforcement (2025\u20132028) creates mandatory compliance demand "
    "for 80,000+ defense contractors, most of whom lack in-house SOC capability.",
    size=9.5
)

# ============================================================
# BUSINESS MODEL
# ============================================================
add_section_header("Business Model")

add_bullet("Appliance + annual subscription (software updates, threat intelligence, model updates)", bold_prefix="Hardware tiers: ")
add_bullet("Small ($X,XXX/yr) \u2264 25 endpoints  |  Medium ($XX,XXX/yr) \u2264 150 endpoints  |  Virtual (enterprise pricing)", bold_prefix="Pricing: ")
add_bullet("85%+ gross margin on subscription revenue after initial hardware COGS", bold_prefix="Margins: ")
add_bullet("Compliance automation reduces customer audit prep from weeks to hours \u2014 strong retention driver", bold_prefix="Retention: ")

# ============================================================
# COMPETITIVE ADVANTAGE
# ============================================================
add_section_header("Why Kahu Wins")
add_bullet("No data leaves premises \u2014 a hard requirement for CMMC, ITAR, HIPAA, and CUI workloads that cloud SIEM cannot meet", bold_prefix="Data sovereignty: ")
add_bullet("Compliance evidence generated automatically as a byproduct of security operations, not a separate workflow", bold_prefix="Evidence as byproduct: ")
add_bullet("One appliance replaces SIEM + SOAR + GRC tool + vulnerability scanner \u2014 dramatic cost savings for SMBs", bold_prefix="Consolidation: ")
add_bullet("No cloud dependencies, no per-query AI costs, no surprise egress fees", bold_prefix="Predictable cost: ")

# ============================================================
# TRACTION & MILESTONES
# ============================================================
add_section_header("Traction & Roadmap")
add_bullet("Core platform built: triage pipeline, AI investigation, compliance engine, vulnerability scanning, web dashboard")
add_bullet("Docker-based deployment operational with GPU-accelerated local inference")
add_bullet("Target: pilot deployments with 3\u20135 defense contractors in Hawaii (Q4 2026)")
add_bullet("CMMC 2.0 enforcement timeline creates urgency \u2014 first-mover advantage in on-prem AI SOC")

# ============================================================
# TEAM (placeholder)
# ============================================================
add_section_header("Team")
add_body(
    "ComplyHI is based in Honolulu, Hawaii, at the intersection of the Pacific defense community "
    "and the state's growing cybersecurity ecosystem. The team brings deep expertise in compliance "
    "frameworks, security operations, and AI/ML engineering.",
    size=9.5
)

# ============================================================
# FOOTER / CONTACT
# ============================================================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="1A1A2E"/>'
    f'</w:pBdr>'
)
p._p.get_or_add_pPr().append(pBdr)
run = p.add_run("KAHU by ComplyHI")
run.font.size = Pt(9)
run.font.bold = True
run.font.color.rgb = BRAND_DARK
run.font.name = "Calibri"
run = p.add_run("  |  On-Premises AI Security Operations  |  complyhi.com  |  info@complyhi.com")
run.font.size = Pt(8)
run.font.color.rgb = GRAY
run.font.name = "Calibri"

# SAVE
out_path = os.path.join(os.path.dirname(__file__), "Kahu_Investor_OneSheet.docx")
doc.save(out_path)
print(f"Saved to: {out_path}")
